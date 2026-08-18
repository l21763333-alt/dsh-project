"""Word(.docx) 解析器：段落 + 表格纯文本提取（MVP 唯一支持的简历格式）。"""
from app.ai.extractors.base import DocumentExtractor, ExtractResult, register_extractor


@register_extractor
class DocxExtractor(DocumentExtractor):
    """python-docx 提取正文与表格文本。"""

    supported_ext = (".docx",)

    def extract(self, file_path: str) -> ExtractResult:
        from docx import Document  # 延迟导入：仅 .docx 解析时才加载

        doc = Document(file_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)
        return ExtractResult(text="\n".join(parts))
