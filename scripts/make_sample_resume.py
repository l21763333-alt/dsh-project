"""生成一份样例 Word 简历（.docx），用于本地端到端验证。

用法：.venv/bin/python scripts/make_sample_resume.py [输出路径]
默认输出：samples/sample_resume.docx
"""
import sys
from pathlib import Path

from docx import Document


def build(path: str) -> None:
    doc = Document()
    doc.add_heading("个人简历", level=0)

    doc.add_heading("基本信息", level=1)
    doc.add_paragraph("姓名：张伟")
    doc.add_paragraph("电话：13800138000")
    doc.add_paragraph("邮箱：zhangwei@example.com")
    doc.add_paragraph("性别：男")
    doc.add_paragraph("出生日期：1996-03-15")

    doc.add_heading("教育背景", level=1)
    doc.add_paragraph("2014.09 - 2018.06  华南理工大学  计算机科学与技术  本科")

    doc.add_heading("工作经历", level=1)
    doc.add_paragraph("2021.07 - 至今  深圳某互联网公司  高级后端工程师")
    doc.add_paragraph("负责订单交易系统设计与开发，主导微服务拆分，日订单量 50 万+。")
    doc.add_paragraph("2018.07 - 2021.06  广州某软件公司  后端工程师")
    doc.add_paragraph("负责 CRM 系统模块开发，使用 Python/Flask 与 MySQL。")

    doc.add_heading("技能清单", level=1)
    table = doc.add_table(rows=2, cols=4)
    skills = [["语言", "Python", "Java", "Go"], ["框架", "Flask", "Spring", "Django"]]
    for i, row in enumerate(skills):
        for j, cell in enumerate(row):
            table.cell(i, j).text = cell

    doc.add_paragraph("")
    doc.add_paragraph("自我评价：5 年后端开发经验，熟悉高并发系统设计与微服务架构，责任心强。")

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"已生成样例简历: {out.resolve()}")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "samples/sample_resume.docx")
