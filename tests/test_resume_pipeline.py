"""简历上传/解析管线集成测试（FakeLLM，不调用真实 API）。"""
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.core.database import SessionLocal
from app.models.candidate import Candidate
from app.models.resume import Resume
from tests.fakes import DEFAULT_PARSE_RESULT, FakeLLM

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


def _make_docx_bytes(name: str = "张三", phone: str = "13800138000") -> bytes:
    doc = Document()
    doc.add_paragraph(f"姓名：{name}")
    doc.add_paragraph(f"电话：{phone}")
    doc.add_paragraph("3年后端开发经验，熟悉 Python 与 Java")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload(client: TestClient, name="张三", phone="13800138000", fname="张三.docx"):
    return client.post(
        "/api/v1/resumes",
        files={"file": (fname, _make_docx_bytes(name, phone), DOCX_MIME)},
        data={"source": "h5_form"},
    )


def test_upload_rejects_non_docx(client: TestClient) -> None:
    resp = client.post(
        "/api/v1/resumes",
        files={"file": ("resume.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )
    assert resp.status_code == 400
    assert "仅支持 .docx" in resp.json()["detail"]


def test_upload_empty_file_rejected(client: TestClient) -> None:
    resp = client.post(
        "/api/v1/resumes",
        files={"file": ("empty.docx", b"", DOCX_MIME)},
    )
    assert resp.status_code == 400
    assert "为空" in resp.json()["detail"]


def test_upload_and_parse_creates_candidate(
    client: TestClient, monkeypatch
) -> None:
    """上传 → 后台解析完成 → 结构化结果落库 + 候选人创建。"""
    monkeypatch.setattr("app.services.resume_service.get_llm", lambda: FakeLLM())
    resp = _upload(client)
    assert resp.status_code == 201
    resume_id = resp.json()["id"]
    assert resp.json()["parse_status"] == "pending"

    # TestClient 会在响应后同步执行后台任务，直接查详情
    detail = client.get(f"/api/v1/resumes/{resume_id}").json()
    assert detail["parse_status"] == "done"
    assert detail["parsed_json"]["name"] == "张三"

    db = SessionLocal()
    try:
        cand = db.query(Candidate).filter(Candidate.phone == "13800138000").first()
        assert cand is not None
        assert cand.name == "张三"
        assert cand.education == "本科"
        assert "Python" in (cand.skills or [])
        assert cand.resume_text and "3年后端开发经验" in cand.resume_text
        resume = db.get(Resume, resume_id)
        assert resume is not None and resume.candidate_id == cand.id
    finally:
        db.close()


def test_parse_failure_records_error(client: TestClient, monkeypatch) -> None:
    """LLM 输出异常时解析状态置 failed 并记录错误信息。"""
    from app.ai.llm import LLMOutputError

    monkeypatch.setattr(
        "app.services.resume_service.get_llm",
        lambda: FakeLLM(error=LLMOutputError("模型输出无法解析")),
    )
    resp = _upload(client)
    resume_id = resp.json()["id"]
    detail = client.get(f"/api/v1/resumes/{resume_id}").json()
    assert detail["parse_status"] == "failed"
    assert "模型输出无法解析" in detail["parse_error"]


def test_retry_parse_after_failure(client: TestClient, monkeypatch) -> None:
    """失败后可手动重试解析，成功后状态转 done。"""
    from app.ai.llm import LLMOutputError

    monkeypatch.setattr(
        "app.services.resume_service.get_llm",
        lambda: FakeLLM(error=LLMOutputError("boom")),
    )
    resp = _upload(client)
    resume_id = resp.json()["id"]
    assert client.get(f"/api/v1/resumes/{resume_id}").json()["parse_status"] == "failed"

    monkeypatch.setattr("app.services.resume_service.get_llm", lambda: FakeLLM())
    retried = client.post(f"/api/v1/resumes/{resume_id}/parse").json()
    assert retried["parse_status"] == "done"


def test_dedup_same_phone_creates_one_candidate(
    client: TestClient, monkeypatch
) -> None:
    """同一手机号的两份简历 → 只生成一个候选人档案。"""
    monkeypatch.setattr("app.services.resume_service.get_llm", lambda: FakeLLM())
    r1 = _upload(client)
    r2 = _upload(client)
    assert r1.status_code == 201 and r2.status_code == 201

    db = SessionLocal()
    try:
        assert db.query(Candidate).count() == 1
        resumes = db.query(Resume).all()
        assert len(resumes) == 2
        candidate_id = resumes[0].candidate_id
        assert all(r.candidate_id == candidate_id for r in resumes)
    finally:
        db.close()


def test_h5_collect_endpoint(client: TestClient, monkeypatch) -> None:
    """H5 在线表单收集入口可提交并解析简历。"""
    result = dict(DEFAULT_PARSE_RESULT)
    result["name"] = "李四"
    result["phone"] = "13900139000"
    monkeypatch.setattr(
        "app.services.resume_service.get_llm", lambda: FakeLLM(result=result)
    )
    resp = client.post(
        "/api/v1/collect/resume",
        files={
            "file": (
                "李四.docx",
                _make_docx_bytes(name="李四", phone="13900139000"),
                DOCX_MIME,
            )
        },
    )
    assert resp.status_code == 201
    detail = client.get(f"/api/v1/resumes/{resp.json()['id']}").json()
    assert detail["parse_status"] == "done"
    assert detail["parsed_json"]["name"] == "李四"
