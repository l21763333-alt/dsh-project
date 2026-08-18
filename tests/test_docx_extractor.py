"""Word(.docx) 解析器测试：用 python-docx 构造真实文档验证提取。"""
from docx import Document

from app.ai.extractors.base import get_extractor
from app.ai.extractors.docx import DocxExtractor


def _make_docx(path, paragraphs=None, table=None) -> None:
    doc = Document()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    if table:
        t = doc.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
    doc.save(str(path))


def test_docx_extractor_registered() -> None:
    """注册表能按扩展名取到 .docx 解析器。"""
    extractor = get_extractor(".docx")
    assert extractor is not None
    assert isinstance(extractor, DocxExtractor)


def test_docx_extracts_paragraphs(tmp_path) -> None:
    docx_path = tmp_path / "resume.docx"
    _make_docx(docx_path, paragraphs=["张三", "电话：13800138000", "3 年后端开发经验"])
    result = DocxExtractor().extract(str(docx_path))
    assert "张三" in result.text
    assert "13800138000" in result.text
    assert "3 年后端开发经验" in result.text


def test_docx_extracts_tables(tmp_path) -> None:
    """表格按行合并为文本（技能表等）。"""
    docx_path = tmp_path / "resume_table.docx"
    _make_docx(
        docx_path,
        paragraphs=["个人信息"],
        table=[["技能", "Python", "Java"], ["语言", "中文", "英文"]],
    )
    result = DocxExtractor().extract(str(docx_path))
    assert "技能 | Python | Java" in result.text
    assert "语言 | 中文 | 英文" in result.text


def test_docx_empty_table_row_skipped(tmp_path) -> None:
    """空行/空单元格不产生噪音行。"""
    docx_path = tmp_path / "resume_empty.docx"
    _make_docx(docx_path, table=[["", "", ""]])
    result = DocxExtractor().extract(str(docx_path))
    assert result.text.strip() == ""
    assert result.images == []
